"""Agent 1 도구: 파일 유형 감지, 유형별 파싱, 품질 평가, NER, 중복 판정.

각 도구는 (입력→출력)이 명확하고, 실패 시 예외를 던져 그래프가 fallback 하도록 한다.
"""
from __future__ import annotations
import base64
import hashlib
import json
import mimetypes
import os
import re

from config import CFG
from llm import chat_json

EXT_MAP = {
    ".pdf": "pdf", ".txt": "txt", ".md": "txt", ".csv": "txt",
    ".hwp": "hwp", ".hwpx": "hwp",
    ".doc": "doc", ".docx": "doc",
    ".jpg": "image", ".jpeg": "image", ".png": "image",
    ".mp4": "audio", ".m4a": "audio", ".wav": "audio",
}

# 유형별 파싱 전략 우선순위 (앞에서부터 시도, 실패/저품질 시 다음으로 fallback)
STRATEGIES = {
    "pdf": ["pdf_text", "vlm_ocr_pdf"],          # 텍스트레이어 우선 → 실패/저품질 시 VLM OCR
    "hwp": ["hwp_libreoffice", "hwp_ole"],
    "doc": ["doc_libreoffice", "doc_python"],
    "txt": ["txt"],
    "image": ["ocr_image"],                       # OCR_BACKEND(vlm/surya/paddleocr)로 분기
    "audio": ["stt"],                             # STT_BACKEND(faster_whisper)로 분기
    "unknown": ["txt"],
}

# 압수 C드라이브에서 제외할 시스템/캐시 디렉터리 (소문자, 경로 조각 매칭)
EXCLUDE_DIRS = {
    "windows", "program files", "program files (x86)", "programdata",
    "$recycle.bin", "system volume information", "appdata", "perflogs",
    "$windows.~bt", "$windows.~ws", "recovery", "msocache",
}

# 수집 대상 확장자 (그 외는 ingest에서 skip)
INGEST_EXTS = set(EXT_MAP.keys())


def is_excluded(path: str) -> bool:
    """경로에 시스템 디렉터리가 포함되면 True."""
    parts = [p.lower() for p in path.replace("\\", "/").split("/")]
    return any(p in EXCLUDE_DIRS for p in parts)


def detect_file_type(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    return EXT_MAP.get(ext, "unknown")


# ---------------- 파싱 전략들 ----------------
def parse_txt(path: str) -> str:
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_pdf_text(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join((p.extract_text() or "") for p in reader.pages)


def _extract_upstage_text(payload: dict) -> str:
    """Upstage Document Parse 응답에서 사용 가능한 텍스트 표현을 꺼낸다."""
    content = payload.get("content") if isinstance(payload, dict) else None
    if isinstance(content, dict):
        for key in ("markdown", "text", "html"):
            val = content.get(key)
            if isinstance(val, str) and val.strip():
                return val
    elements = payload.get("elements") if isinstance(payload, dict) else None
    chunks: list[str] = []
    for el in elements if isinstance(elements, list) else []:
        c = el.get("content") if isinstance(el, dict) else None
        if isinstance(c, dict):
            for key in ("markdown", "text", "html"):
                val = c.get(key)
                if isinstance(val, str) and val.strip():
                    chunks.append(val.strip())
                    break
        elif isinstance(c, str) and c.strip():
            chunks.append(c.strip())
    return "\n".join(chunks)


def upstage_document_parse(path: str) -> str:
    """Upstage Document Parse로 PDF/이미지/HWP/DOC 계열 문서를 구조화 텍스트로 변환."""
    if not CFG.upstage_api_key:
        raise RuntimeError("UPSTAGE_API_KEY 미설정")
    import requests

    mime = mimetypes.guess_type(path)[0] or "application/octet-stream"
    data = {
        "model": CFG.upstage_document_model,
        "output_formats": json.dumps([CFG.upstage_document_output_format], ensure_ascii=False),
    }
    if CFG.upstage_document_ocr:
        data["ocr"] = CFG.upstage_document_ocr
    headers = {"Authorization": f"Bearer {CFG.upstage_api_key}"}
    with open(path, "rb") as f:
        files = {"document": (os.path.basename(path), f, mime)}
        resp = requests.post(
            CFG.upstage_document_endpoint,
            headers=headers,
            files=files,
            data=data,
            timeout=CFG.upstage_timeout_sec,
        )
    if resp.status_code >= 400:
        raise RuntimeError(f"Upstage Document Parse 실패 {resp.status_code}: {resp.text[:300]}")
    try:
        payload = resp.json()
    except Exception as e:
        raise RuntimeError(f"Upstage Document Parse JSON 파싱 실패: {e}") from e
    text = _extract_upstage_text(payload)
    if not text.strip():
        raise RuntimeError("Upstage Document Parse 결과 텍스트 없음")
    return text


def vlm_ocr_pdf(path: str) -> str:
    """pypdfium2 로 페이지 렌더 → 비전 모델 OCR. (옵션 의존성)"""
    import pypdfium2 as pdfium  # 미설치 시 예외 → fallback 종료
    from llm import vlm_client, vlm_model_name
    pdf = pdfium.PdfDocument(path)
    texts = []
    for i in range(min(len(pdf), 5)):  # 프로토타입: 최대 5p
        img = pdf[i].render(scale=2).to_pil()
        import io
        buf = io.BytesIO(); img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        r = vlm_client().chat.completions.create(
            model=vlm_model_name(), temperature=0,
            messages=[{"role": "user", "content": [
                {"type": "text", "text": "이 이미지의 모든 한국어 텍스트를 그대로 추출하라."},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            ]}],
        )
        texts.append(r.choices[0].message.content or "")
    return "\n".join(texts)


def vlm_ocr_image(path: str) -> str:
    from llm import vlm_client, vlm_model_name
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    ext = os.path.splitext(path)[1].lstrip(".") or "png"
    r = vlm_client().chat.completions.create(
        model=vlm_model_name(), temperature=0,
        messages=[{"role": "user", "content": [
            {"type": "text", "text": "이미지 속 한국어 텍스트를 모두 추출하라. 없으면 빈 문자열."},
            {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{b64}"}},
        ]}],
    )
    return r.choices[0].message.content or ""


def vlm_caption_image(path: str) -> str:
    """이미지에 추출할 텍스트가 없을 때 사진 내용을 한국어로 간결히 설명(캡션).
    전체 파일 색인을 위해 사용 — OCR 결과가 빈 경우의 보완."""
    from llm import vlm_client, vlm_model_name
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    ext = os.path.splitext(path)[1].lstrip(".") or "png"
    r = vlm_client().chat.completions.create(
        model=vlm_model_name(), temperature=0,
        messages=[{"role": "user", "content": [
            {"type": "text", "text": "이 이미지를 한국어 한두 문장으로 객관적으로 설명하라. "
                                     "보이는 사람·사물·장면 위주로. 추측성 서술은 피하라."},
            {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{b64}"}},
        ]}],
    )
    return (r.choices[0].message.content or "").strip()


def hwp_ole(path: str) -> str:
    """olefile 기반 HWP 본문 추출(간이). 실패 시 예외 → libreoffice fallback."""
    import olefile
    import zlib
    if not olefile.isOleFile(path):
        raise ValueError("OLE 형식 아님(hwpx일 수 있음)")
    ole = olefile.OleFileIO(path)
    texts = []
    for entry in ole.listdir():
        if entry and entry[0] == "BodyText":
            data = ole.openstream(entry).read()
            try:
                data = zlib.decompress(data, -15)
            except Exception:
                pass
            # UTF-16LE 추정 후 한글/공백만 보존
            try:
                s = data.decode("utf-16le", errors="ignore")
            except Exception:
                s = ""
            s = "".join(ch for ch in s if ch == " " or ch == "\n" or "가" <= ch <= "힣"
                        or ch.isalnum())
            if s.strip():
                texts.append(s)
    ole.close()
    if not texts:
        raise ValueError("HWP 본문 추출 실패")
    return "\n".join(texts)


def hwp_libreoffice(path: str) -> str:
    """LibreOffice(soffice) 가 있으면 txt 변환."""
    import shutil
    import subprocess
    import tempfile
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice:
        raise FileNotFoundError("LibreOffice(soffice) 미설치")
    out = tempfile.mkdtemp()
    subprocess.run([soffice, "--headless", "--convert-to", "txt:Text",
                    "--outdir", out, path], check=True, timeout=120)
    base = os.path.splitext(os.path.basename(path))[0] + ".txt"
    return parse_txt(os.path.join(out, base))


def doc_python(path: str) -> str:
    """python-docx 기반 .docx 추출. .doc(바이너리)는 실패 → libreoffice fallback."""
    import docx  # python-docx
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


def doc_libreoffice(path: str) -> str:
    """LibreOffice로 doc/docx → txt 변환 (hwp 변환 로직 재사용)."""
    return hwp_libreoffice(path)


def _extract_clova_text(payload: dict) -> str:
    """CLOVA Speech 응답에서 전사 본문을 추출한다."""
    for key in ("text", "fullText", "transcript"):
        val = payload.get(key) if isinstance(payload, dict) else None
        if isinstance(val, str) and val.strip():
            return val.strip()
    segments = payload.get("segments") if isinstance(payload, dict) else None
    texts = []
    for seg in segments if isinstance(segments, list) else []:
        if isinstance(seg, dict) and isinstance(seg.get("text"), str):
            texts.append(seg["text"].strip())
    if texts:
        return "\n".join(t for t in texts if t)
    speakers = payload.get("speakers") if isinstance(payload, dict) else None
    for spk in speakers if isinstance(speakers, list) else []:
        if isinstance(spk, dict) and isinstance(spk.get("text"), str):
            texts.append(spk["text"].strip())
    return "\n".join(t for t in texts if t)


def _clova_headers(json_body: bool = True) -> dict:
    if not CFG.clova_secret_key:
        raise RuntimeError("CLOVA_SECRET_KEY 미설정")
    headers = {
        "Accept": "application/json;UTF-8",
        "X-CLOVASPEECH-API-KEY": CFG.clova_secret_key,
    }
    if json_body:
        headers["Content-Type"] = "application/json;UTF-8"
    return headers


def _clova_request_body() -> dict:
    return {
        "language": CFG.clova_language,
        "completion": CFG.clova_completion_mode,
        "wordAlignment": True,
        "fullText": True,
        "diarization": {"enable": True},
    }


def _upload_to_ncloud_object_storage(path: str) -> str:
    if not (CFG.ncloud_access_key and CFG.ncloud_secret_key and CFG.ncloud_bucket_name):
        raise RuntimeError("NCLOUD Object Storage 설정 미완성")
    import boto3

    prefix = CFG.ncloud_clova_input_prefix.strip("/")
    key = f"{prefix}/{os.path.basename(path)}" if prefix else os.path.basename(path)
    s3 = boto3.client(
        "s3",
        endpoint_url=CFG.ncloud_object_storage_endpoint,
        aws_access_key_id=CFG.ncloud_access_key,
        aws_secret_access_key=CFG.ncloud_secret_key,
        region_name=CFG.ncloud_region,
    )
    s3.upload_file(path, CFG.ncloud_bucket_name, key)
    return key


def _clova_object_storage(path: str) -> dict:
    import requests

    data_key = _upload_to_ncloud_object_storage(path)
    body = _clova_request_body()
    body["dataKey"] = data_key
    url = CFG.clova_invoke_url.rstrip("/") + "/recognizer/object-storage"
    resp = requests.post(
        url,
        headers=_clova_headers(json_body=True),
        data=json.dumps(body, ensure_ascii=False).encode("UTF-8"),
        timeout=CFG.clova_timeout_sec,
    )
    if resp.status_code >= 400:
        raise RuntimeError(f"CLOVA object-storage 실패 {resp.status_code}: {resp.text[:300]}")
    return resp.json()


def _clova_upload(path: str) -> dict:
    import requests

    body = _clova_request_body()
    url = CFG.clova_invoke_url.rstrip("/") + "/recognizer/upload"
    with open(path, "rb") as f:
        files = {
            "media": (os.path.basename(path), f, mimetypes.guess_type(path)[0] or "application/octet-stream"),
            "params": (
                None,
                json.dumps(body, ensure_ascii=False).encode("UTF-8"),
                "application/json",
            ),
        }
        resp = requests.post(
            url,
            headers=_clova_headers(json_body=False),
            files=files,
            timeout=CFG.clova_timeout_sec,
        )
    if resp.status_code >= 400:
        raise RuntimeError(f"CLOVA upload 실패 {resp.status_code}: {resp.text[:300]}")
    return resp.json()


def clova_speech(path: str) -> str:
    """CLOVA Speech 장문 인식으로 mp4/m4a/wav 파일을 전사."""
    if not CFG.clova_invoke_url:
        raise RuntimeError("CLOVA_INVOKE_URL 미설정")
    if not CFG.clova_secret_key:
        raise RuntimeError("CLOVA_SECRET_KEY 미설정")
    if CFG.clova_use_object_storage and CFG.ncloud_bucket_name and CFG.ncloud_access_key and CFG.ncloud_secret_key:
        payload = _clova_object_storage(path)
    else:
        payload = _clova_upload(path)
    text = _extract_clova_text(payload)
    if not text.strip():
        raise RuntimeError(f"CLOVA Speech 결과 텍스트 없음: {str(payload)[:300]}")
    return text


def whisper(path: str) -> str:
    """이전 전략명 호환용. 실제 구현은 CLOVA Speech를 사용한다(레거시)."""
    return clova_speech(path)


# ---------------- 로컬 STT (faster-whisper) ----------------
_whisper_model = None
_paddle_model = None


def faster_whisper_stt(path: str) -> str:
    """로컬 faster-whisper로 음성(mp4/m4a/wav) 전사. m4a는 내부 PyAV로 직접 디코드."""
    global _whisper_model
    if _whisper_model is None:
        from faster_whisper import WhisperModel  # 지연 임포트
        _whisper_model = WhisperModel(CFG.whisper_model, device=CFG.stt_device,
                                      compute_type=CFG.stt_compute_type)
    lang = (CFG.clova_language or "ko-KR").split("-")[0]  # "ko"
    segments, _info = _whisper_model.transcribe(path, language=lang, vad_filter=True)
    text = "\n".join(seg.text.strip() for seg in segments).strip()
    if not text:
        raise RuntimeError("faster-whisper 전사 결과 없음")
    return text


def stt(path: str) -> str:
    """STT 디스패처: STT_BACKEND=faster_whisper(기본) | clova(레거시)."""
    if CFG.stt_backend == "clova":
        return clova_speech(path)
    return faster_whisper_stt(path)


# ---------------- 로컬 OCR (vlm 기본 / surya / paddleocr) ----------------
def _surya_ocr(path: str) -> str:
    """Surya OCR(베스트 에포트). API 버전차로 실패 시 호출부가 VLM으로 폴백."""
    from PIL import Image
    from surya.recognition import RecognitionPredictor
    from surya.detection import DetectionPredictor
    img = Image.open(path).convert("RGB")
    rec, det = RecognitionPredictor(), DetectionPredictor()
    preds = rec([img], det_predictor=det)
    lines = getattr(preds[0], "text_lines", [])
    return "\n".join(getattr(ln, "text", "") for ln in lines).strip()


def _paddle_ocr(path: str) -> str:
    """PaddleOCR(베스트 에포트). 실패 시 호출부가 VLM으로 폴백."""
    global _paddle_model
    if _paddle_model is None:
        from paddleocr import PaddleOCR
        _paddle_model = PaddleOCR(use_angle_cls=True, lang="korean")
    res = _paddle_model.ocr(path, cls=True)
    out = []
    for page in res or []:
        for line in page or []:
            try:
                out.append(line[1][0])
            except Exception:
                pass
    return "\n".join(out).strip()


def ocr_image(path: str) -> str:
    """이미지 OCR 디스패처: OCR_BACKEND=vlm(기본·견고) | surya | paddleocr. 실패 시 VLM 폴백."""
    b = CFG.ocr_backend
    if b == "surya":
        try:
            t = _surya_ocr(path)
            if t:
                return t
        except Exception:
            pass
    elif b == "paddleocr":
        try:
            t = _paddle_ocr(path)
            if t:
                return t
        except Exception:
            pass
    return vlm_ocr_image(path)


PARSERS = {
    # 로컬 전략 (STRATEGIES에서 사용)
    "pdf_text": parse_pdf_text, "vlm_ocr_pdf": vlm_ocr_pdf,
    "hwp_libreoffice": hwp_libreoffice, "hwp_ole": hwp_ole,
    "doc_libreoffice": doc_libreoffice, "doc_python": doc_python,
    "txt": parse_txt,
    "ocr_image": ocr_image, "vlm_ocr_image": vlm_ocr_image,
    "stt": stt, "faster_whisper": faster_whisper_stt,
    # 레거시(클라우드) — STRATEGIES 미사용, 필요 시 명시 호출만
    "upstage_document_parse": upstage_document_parse,
    "clova_speech": clova_speech, "whisper": whisper,
}


# ---------------- 품질 평가 ----------------
def assess_extraction_quality(text: str) -> float:
    """0~1 점수. 길이/인쇄가능비율/치환문자(�) 기반 휴리스틱."""
    if not text or not text.strip():
        return 0.0
    n = len(text)
    repl = text.count("�")
    printable = sum(1 for c in text if c.isprintable() or c in "\n\t")
    length_score = min(n / 200.0, 1.0)          # 200자 이상이면 충분
    clean_score = max(0.0, 1.0 - repl / max(n, 1) * 5)
    print_score = printable / max(n, 1)
    return round(0.4 * length_score + 0.3 * clean_score + 0.3 * print_score, 3)


# ---------------- 중복 판정 ----------------
def text_fingerprint(text: str) -> str:
    norm = re.sub(r"\s+", " ", (text or "").strip().lower())
    return hashlib.sha1(norm.encode("utf-8")).hexdigest()


# ---------------- NER ----------------
_NER_SYS = (
    "너는 한국 디지털 포렌식 수사용 개체명 인식기다. 주어진 문서에서 다음 5개 유형의 "
    "증거 개체를 빠짐없이(재현율 우선) 추출하라: 인물, 날짜, 금액, 계좌, 장소. "
    "각 개체에 대해 entity_text(원문 표기), entity_type(위 5개 중 하나), "
    "context(개체가 등장한 짧은 원문 문맥)를 포함하라. "
    'JSON 배열로만 출력: [{"entity_text":..,"entity_type":..,"context":..}, ...]'
)


def _as_entity_list(data) -> list[dict]:
    if isinstance(data, dict):
        data = data.get("entities") or data.get("results") or []
    out = []
    for e in data if isinstance(data, list) else []:
        if isinstance(e, dict) and e.get("entity_text") and e.get("entity_type"):
            out.append({
                "entity_text": str(e["entity_text"]).strip(),
                "entity_type": str(e["entity_type"]).strip(),
                "context": str(e.get("context", "")).strip(),
            })
    return out


def run_ner(text: str) -> list[dict]:
    """NER. self_consistency_n>1이면 K회 실행 후 '합집합'(누락 최소화=Recall 우선)."""
    if not text.strip():
        return []
    snippet = text[:6000]  # 프로토타입 토큰 제한
    n = max(1, CFG.self_consistency_n)
    if n == 1:
        return _as_entity_list(chat_json(_NER_SYS, f"문서:\n{snippet}"))
    merged: dict[tuple, dict] = {}
    for _ in range(n):  # 다양성 위해 시드 미적용·온도 상향
        for e in _as_entity_list(chat_json(_NER_SYS, f"문서:\n{snippet}",
                                            temperature=0.4, use_seed=False)):
            merged.setdefault((e["entity_text"], e["entity_type"]), e)
    return list(merged.values())
